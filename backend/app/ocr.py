"""
OCR Invoice Extraction Module

Extracts structured data from invoice images/PDFs using Tesseract OCR.
Parses vendor GSTIN, invoice details, line items, and GST amounts.
"""

import re
import json
from datetime import datetime
from typing import Dict, List, Optional
from pathlib import Path

try:
    import pytesseract
    from PIL import Image
    import pdf2image
    from docx import Document
    import pandas as pd
except ImportError:
    print("Warning: OCR dependencies not installed. Run: pip install pytesseract pdf2image Pillow python-docx pandas")
    pytesseract = None
    Image = None
    pdf2image = None
    Document = None
    pd = None


class InvoiceOCR:
    """Extract and parse invoice data from images and PDFs"""

    # GSTIN format: 2 digits + 5 letters + 4 digits + 1 letter + 1 alphanumeric + Z + 1 alphanumeric
    GSTIN_PATTERN = r'\b\d{2}[A-Z]{5}\d{4}[A-Z]{1}[A-Z\d]{1}[Z]{1}[A-Z\d]{1}\b'

    # Common invoice number patterns
    INVOICE_NUM_PATTERNS = [
        r'Invoice\s*(?:No|Number|#)?\.?\s*:?\s*([A-Z0-9\-/]+)',
        r'Bill\s*(?:No|Number|#)?\.?\s*:?\s*([A-Z0-9\-/]+)',
        r'Tax\s*Invoice\s*:?\s*([A-Z0-9\-/]+)',
        r'INV[:\s\-]*([A-Z0-9\-/]+)',
        r'(?:Receipt|Ref)\s*(?:No|Number|#)?\.?\s*:?\s*([A-Z0-9\-/]+)',
        # Pattern for P/YYYY-YY/NNNN format (common in government/utility bills)
        r'([A-Z]/\d{4}-\d{2}/\d+)',
        # Pattern for standalone alphanumeric codes after "Invoice No."
        r'Invoice\s+No\.?\s*:?\s+([A-Z0-9/\-]+)',
    ]

    # Date patterns (DD/MM/YYYY, DD-MM-YYYY, etc.)
    DATE_PATTERNS = [
        r'\b(\d{1,2})[/-](\d{1,2})[/-](\d{4})\b',  # DD/MM/YYYY or DD-MM-YYYY
        r'\b(\d{4})[/-](\d{1,2})[/-](\d{1,2})\b',  # YYYY-MM-DD
    ]

    def __init__(self):
        if pytesseract is None:
            raise ImportError(
                "OCR dependencies not installed. Install with:\n"
                "pip install pytesseract pdf2image Pillow\n"
                "Also install Tesseract: brew install tesseract (Mac) or apt-get install tesseract-ocr (Linux)"
            )

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from invoice file using OCR or text extraction

        Args:
            file_path: Path to invoice file (PDF, JPG, PNG, DOCX, XLSX)

        Returns:
            Extracted text content
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Invoice file not found: {file_path}")

        # Handle PDF files
        if file_path.suffix.lower() == '.pdf':
            return self._extract_from_pdf(file_path)

        # Handle image files
        elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return self._extract_from_image(file_path)
        
        # Handle Word documents
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        
        # Handle Excel files
        elif file_path.suffix.lower() in ['.xlsx', '.xls']:
            return self._extract_from_excel(file_path)

        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

    def _extract_from_image(self, image_path: Path) -> str:
        """Extract text from image file"""
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            raise RuntimeError(f"OCR failed for image {image_path}: {str(e)}")

    def _extract_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF file (converts to images first)"""
        try:
            # Convert PDF pages to images
            images = pdf2image.convert_from_path(str(pdf_path))

            # Extract text from each page
            full_text = ""
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image)
                full_text += f"\n--- Page {i+1} ---\n{page_text}"

            return full_text
        except pdf2image.exceptions.PDFInfoNotInstalledError:
            raise RuntimeError(
                f"Poppler not installed. Install with: brew install poppler (Mac) or apt-get install poppler-utils (Linux)"
            )
        except Exception as e:
            raise RuntimeError(f"OCR failed for PDF {pdf_path}: {str(e)}")

    def _extract_from_docx(self, docx_path: Path) -> str:
        """Extract text from Word document"""
        try:
            if Document is None:
                raise RuntimeError("python-docx not installed")
            
            doc = Document(str(docx_path))
            full_text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + "\t"
                    full_text += "\n"
            
            return full_text
        except Exception as e:
            raise RuntimeError(f"Failed to extract from Word document {docx_path}: {str(e)}")

    def _extract_from_excel(self, excel_path: Path) -> str:
        """Extract text from Excel file"""
        try:
            if pd is None:
                raise RuntimeError("pandas not installed")
            
            # Read all sheets
            excel_file = pd.ExcelFile(str(excel_path))
            full_text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                full_text += f"\n--- Sheet: {sheet_name} ---\n"
                full_text += df.to_string(index=False)
                full_text += "\n"
            
            return full_text
        except Exception as e:
            raise RuntimeError(f"Failed to extract from Excel file {excel_path}: {str(e)}")

    def extract_gstin(self, text: str) -> Optional[str]:
        """
        Extract GSTIN from invoice text

        Args:
            text: OCR extracted text

        Returns:
            GSTIN if found, None otherwise
        """
        matches = re.findall(self.GSTIN_PATTERN, text)

        if matches:
            # Return first match (usually vendor GSTIN appears first)
            return matches[0]

        return None

    def extract_invoice_number(self, text: str) -> Optional[str]:
        """Extract invoice number from text"""
        for pattern in self.INVOICE_NUM_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()

        return None

    def extract_date(self, text: str) -> Optional[str]:
        """
        Extract invoice date from text

        Returns date in YYYY-MM-DD format
        """
        for pattern in self.DATE_PATTERNS:
            match = re.search(pattern, text)
            if match:
                groups = match.groups()

                # Handle DD/MM/YYYY format
                if len(groups[0]) <= 2:
                    day, month, year = groups
                    try:
                        date_obj = datetime(int(year), int(month), int(day))
                        return date_obj.strftime('%Y-%m-%d')
                    except ValueError:
                        continue

                # Handle YYYY-MM-DD format
                else:
                    year, month, day = groups
                    try:
                        date_obj = datetime(int(year), int(month), int(day))
                        return date_obj.strftime('%Y-%m-%d')
                    except ValueError:
                        continue

        return None

    def extract_amounts(self, text: str) -> Dict[str, float]:
        """
        Extract GST amounts (CGST, SGST, IGST) and total from text

        Returns:
            Dictionary with cgst, sgst, igst, total_amount
        """
        amounts = {
            'cgst': 0.0,
            'sgst': 0.0,
            'igst': 0.0,
            'total_amount': 0.0
        }

        # Pattern to match amounts (handles Indian number format with commas)
        amount_pattern = r'(?:Rs\.?|INR)?\s*([0-9,]+(?:\.\d{2})?)'

        # Extract CGST
        cgst_match = re.search(r'CGST\s*(?:@\s*\d+\.?\d*%?)?\s*:?\s*' + amount_pattern, text, re.IGNORECASE)
        if cgst_match:
            amounts['cgst'] = self._parse_amount(cgst_match.group(1))

        # Extract SGST
        sgst_match = re.search(r'SGST\s*(?:@\s*\d+\.?\d*%?)?\s*:?\s*' + amount_pattern, text, re.IGNORECASE)
        if sgst_match:
            amounts['sgst'] = self._parse_amount(sgst_match.group(1))

        # Extract IGST
        igst_match = re.search(r'IGST\s*(?:@\s*\d+\.?\d*%?)?\s*:?\s*' + amount_pattern, text, re.IGNORECASE)
        if igst_match:
            amounts['igst'] = self._parse_amount(igst_match.group(1))

        # Extract Total
        total_patterns = [
            r'Total\s*(?:Amount)?\s*:?\s*' + amount_pattern,
            r'Grand\s*Total\s*:?\s*' + amount_pattern,
            r'Net\s*Amount\s*:?\s*' + amount_pattern,
        ]

        for pattern in total_patterns:
            total_match = re.search(pattern, text, re.IGNORECASE)
            if total_match:
                amounts['total_amount'] = self._parse_amount(total_match.group(1))
                break

        return amounts

    def _parse_amount(self, amount_str: str) -> float:
        """Convert amount string to float (handles commas)"""
        # Remove commas and convert to float
        clean_amount = amount_str.replace(',', '')
        try:
            return float(clean_amount)
        except ValueError:
            return 0.0

    def extract_line_items(self, text: str) -> List[Dict]:
        """
        Extract line items from invoice (simplified version)

        This is a basic implementation. For production, you'd need more
        sophisticated parsing based on invoice format/templates.

        Returns:
            List of line items with description, quantity, rate, amount
        """
        line_items = []

        # This is a placeholder - real implementation would need:
        # 1. Table detection (using computer vision or regex patterns)
        # 2. Column identification (description, qty, rate, amount)
        # 3. Row extraction

        # For MVP, we'll return empty list and rely on manual entry
        # or structured data extraction in future versions

        return line_items

    def extract_invoice_data(self, file_path: str) -> Dict:
        """
        Main method to extract all invoice data

        Args:
            file_path: Path to invoice file

        Returns:
            Dictionary with all extracted invoice data
        """
        try:
            # Extract text using OCR
            text = self.extract_text_from_file(file_path)

            # Extract structured data
            gstin = self.extract_gstin(text)
            invoice_number = self.extract_invoice_number(text)
            invoice_date = self.extract_date(text)
            amounts = self.extract_amounts(text)
            line_items = self.extract_line_items(text)

            # Calculate OCR confidence (basic heuristic)
            confidence = self._calculate_confidence(gstin, invoice_number, invoice_date, amounts)

            return {
                'success': True,
                'confidence': confidence,
                'vendor_gstin': gstin,
                'invoice_number': invoice_number,
                'invoice_date': invoice_date,
                'line_items': line_items,
                'total_amount': amounts['total_amount'],
                'cgst': amounts['cgst'],
                'sgst': amounts['sgst'],
                'igst': amounts['igst'],
                'raw_text': text,  # Store raw text for debugging
                'extracted_at': datetime.now().isoformat()
            }
        except Exception as e:
            # Return error dict instead of crashing
            return {
                'success': False,
                'error': str(e),
                'error_type': type(e).__name__,
                'confidence': 0,
                'vendor_gstin': None,
                'invoice_number': None,
                'invoice_date': None,
                'line_items': [],
                'total_amount': 0.0,
                'cgst': 0.0,
                'sgst': 0.0,
                'igst': 0.0,
                'extracted_at': datetime.now().isoformat()
            }

    def _calculate_confidence(self, gstin: Optional[str], invoice_num: Optional[str],
                            invoice_date: Optional[str], amounts: Dict) -> int:
        """
        Calculate OCR confidence score (0-100)

        Higher score = more fields successfully extracted
        """
        score = 0

        # GSTIN found (+30 points)
        if gstin:
            score += 30

        # Invoice number found (+20 points)
        if invoice_num:
            score += 20

        # Date found (+20 points)
        if invoice_date:
            score += 20

        # Total amount found (+20 points)
        if amounts.get('total_amount', 0) > 0:
            score += 20

        # GST amounts found (+10 points)
        if amounts.get('cgst', 0) > 0 or amounts.get('igst', 0) > 0:
            score += 10

        return min(score, 100)  # Cap at 100


# Convenience function for direct use
def extract_invoice_data(file_path: str) -> Dict:
    """
    Extract invoice data from file (convenience function)

    Usage:
        from app.ocr import extract_invoice_data
        result = extract_invoice_data('invoice.pdf')
        print(result)
    """
    ocr = InvoiceOCR()
    return ocr.extract_invoice_data(file_path)


# CLI for testing
if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Usage: python ocr.py <invoice_file>")
        print("Example: python ocr.py ../test_invoices/sample1.pdf")
        sys.exit(1)

    invoice_file = sys.argv[1]

    print(f"Processing invoice: {invoice_file}")
    print("-" * 60)

    try:
        result = extract_invoice_data(invoice_file)

        # Pretty print results
        print(json.dumps(result, indent=2))

        print("\n" + "=" * 60)
        print(f"Extraction Confidence: {result['confidence']}%")

        if result['confidence'] < 50:
            print("\n⚠️  Warning: Low confidence. Consider manual review.")
        elif result['confidence'] < 80:
            print("\n✓ Moderate confidence. Please verify extracted data.")
        else:
            print("\n✓✓ High confidence. Data looks good!")

    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        sys.exit(1)
